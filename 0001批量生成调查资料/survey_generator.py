#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
批量生成调查表工具
功能：读取Shapefile图层数据，使用Word模板批量填充占位符生成调查表
作者：Claude Code
日期：2026-02-06
"""

import os
import re
import sys
from pathlib import Path
from typing import Dict, List, Tuple, Iterator, Any
import warnings

# 忽略geopandas的警告
warnings.filterwarnings('ignore')

try:
    import geopandas as gpd
    from docx import Document
    from tqdm import tqdm
except ImportError as e:
    print(f"错误: 缺少必要的依赖库")
    print(f"请运行: pip install -r requirements.txt")
    print(f"详细信息: {e}")
    sys.exit(1)


class ShapefileReader:
    """Shapefile读取器"""

    def __init__(self, shp_path: str, encoding: str = 'gbk'):
        """
        初始化Shapefile读取器

        Args:
            shp_path: Shapefile文件路径
            encoding: 文件编码，默认gbk
        """
        self.shp_path = shp_path
        self.encoding = encoding
        self.gdf = None
        self._read()

    def _read(self):
        """读取Shapefile"""
        try:
            self.gdf = gpd.read_file(self.shp_path, encoding=self.encoding)
        except Exception as e:
            # 尝试其他编码
            if self.encoding == 'gbk':
                try:
                    self.gdf = gpd.read_file(self.shp_path, encoding='utf-8')
                    self.encoding = 'utf-8'
                except:
                    raise Exception(f"无法读取Shapefile: {e}\n请检查文件路径和编码格式")
            else:
                raise Exception(f"无法读取Shapefile: {e}")

    def get_fields(self) -> List[str]:
        """获取所有字段名"""
        return [col for col in self.gdf.columns if col != 'geometry']

    def get_field_info(self) -> List[Dict[str, Any]]:
        """获取字段详细信息"""
        field_info = []
        for col in self.get_fields():
            dtype = self.gdf[col].dtype
            sample = None
            if len(self.gdf) > 0:
                sample_val = self.gdf[col].iloc[0]
                sample = str(sample_val) if sample_val is not None else None
            field_info.append({
                'name': col,
                'type': str(dtype),
                'sample': sample
            })
        return field_info

    def get_record_count(self) -> int:
        """获取记录数量"""
        return len(self.gdf)

    def get_records(self) -> Iterator[Dict[str, Any]]:
        """返回记录迭代器"""
        for idx, row in self.gdf.iterrows():
            record = {}
            for col in self.get_fields():
                value = row[col]
                # 处理空值
                if value is None or (hasattr(value, '__class__') and value.__class__.__name__ == 'float' and str(value) == 'nan'):
                    record[col] = ''
                else:
                    record[col] = str(value).strip()
            yield record


class TemplateProcessor:
    """Word模板处理器"""

    def __init__(self, template_path: str):
        """
        初始化模板处理器

        Args:
            template_path: Word模板文件路径
        """
        self.template_path = template_path
        self.placeholders = []
        self._detect_placeholders()

    def _detect_placeholders(self):
        """检测模板中的所有占位符"""
        try:
            doc = Document(self.template_path)
            placeholders = set()

            # 扫描段落
            for paragraph in doc.paragraphs:
                matches = re.findall(r'!([A-Za-z0-9_\u4e00-\u9fa5]+)!', paragraph.text)
                placeholders.update(matches)

            # 扫描表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(r'!([A-Za-z0-9_\u4e00-\u9fa5]+)!', cell.text)
                        placeholders.update(matches)

            self.placeholders = list(placeholders)
        except Exception as e:
            raise Exception(f"无法读取Word模板: {e}")

    def get_placeholders(self) -> List[str]:
        """获取所有占位符"""
        return self.placeholders

    def render(self, data: Dict[str, str], output_path: str) -> bool:
        """
        渲染模板并保存

        Args:
            data: 占位符数据字典 {字段名: 值}
            output_path: 输出文件路径

        Returns:
            bool: 是否成功
        """
        try:
            doc = Document(self.template_path)

            # 替换段落中的占位符
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    placeholder = f'!{key}!'
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

            # 替换表格中的占位符
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            placeholder = f'!{key}!'
                            if placeholder in cell.text:
                                cell.text = cell.text.replace(placeholder, str(value))

            # 保存文档
            doc.save(output_path)
            return True

        except Exception as e:
            print(f"  警告: 渲染失败 - {e}")
            return False


class BatchGenerator:
    """批量生成器"""

    def __init__(self, shp_reader: ShapefileReader, template_processor: TemplateProcessor):
        """
        初始化批量生成器

        Args:
            shp_reader: Shapefile读取器
            template_processor: 模板处理器
        """
        self.shp_reader = shp_reader
        self.template_processor = template_processor
        self.filename_counter = {}  # 跟踪文件名使用次数，处理冲突

    def generate_all(self, output_dir: str, naming_field: str) -> Dict[str, Any]:
        """
        批量生成所有文档

        Args:
            output_dir: 输出目录
            naming_field: 用于命名字段

        Returns:
            生成结果统计
        """
        # 创建输出目录
        os.makedirs(output_dir, exist_ok=True)

        results = {
            'success': [],
            'failed': [],
            'total': 0
        }

        # 获取所有记录
        records = list(self.shp_reader.get_records())
        results['total'] = len(records)

        print(f"\n正在生成文档...")

        # 批量生成
        for record in tqdm(records, desc="生成进度"):
            try:
                # 生成基础文件名
                base_filename = self._sanitize_filename(str(record.get(naming_field, 'unnamed')))

                # 获取唯一文件名（处理冲突）
                filename = self._get_unique_filename(base_filename)
                output_path = os.path.join(output_dir, f"{filename}.docx")

                # 渲染文档
                success = self.template_processor.render(record, output_path)

                if success:
                    results['success'].append(filename)
                else:
                    results['failed'].append((filename, '渲染失败'))

            except Exception as e:
                filename = str(record.get(naming_field, 'unknown'))
                results['failed'].append((filename, str(e)))

        return results

    def _sanitize_filename(self, filename: str) -> str:
        """
        清理文件名，移除非法字符

        Args:
            filename: 原始文件名

        Returns:
            清理后的文件名
        """
        # 移除Windows非法字符
        cleaned = re.sub(r'[<>:"/\\|?*]', '_', filename)
        cleaned = cleaned.strip()

        # 限制长度
        if len(cleaned) > 200:
            cleaned = cleaned[:200]

        # 如果为空，使用默认名称
        if not cleaned:
            cleaned = 'unnamed'

        return cleaned

    def _get_unique_filename(self, base_filename: str) -> str:
        """
        获取唯一文件名，如果冲突则添加序号后缀

        Args:
            base_filename: 基础文件名（不含扩展名）

        Returns:
            唯一的文件名
        """
        if base_filename not in self.filename_counter:
            # 第一次使用此文件名
            self.filename_counter[base_filename] = 1
            return base_filename
        else:
            # 文件名冲突，添加序号
            self.filename_counter[base_filename] += 1
            count = self.filename_counter[base_filename]
            return f"{base_filename}_{count}"


class InteractiveCLI:
    """交互式命令行界面"""

    def __init__(self):
        self.current_dir = Path.cwd()

    def print_header(self):
        """打印欢迎信息"""
        print("=" * 60)
        print(" " * 15 + "批量生成调查表工具 v1.0")
        print("=" * 60)
        print()

    def select_shapefile(self) -> str:
        """选择Shapefile文件"""
        print("【步骤 1/7】选择Shapefile文件")
        print(f"当前目录: {self.current_dir}")
        print()

        # 自动查找.shp文件
        shp_files = list(self.current_dir.rglob("*.shp"))

        if shp_files:
            print("找到以下Shapefile文件:")
            for i, shp in enumerate(shp_files, 1):
                print(f"  {i}. {shp.relative_to(self.current_dir)}")
            print()

            choice = input(f"请选择 (1-{len(shp_files)}, 或按Enter输入路径): ").strip()

            if choice.isdigit() and 1 <= int(choice) <= len(shp_files):
                selected = shp_files[int(choice) - 1]
                print(f"✓ 已选择: {selected.relative_to(self.current_dir)}")
                print()
                return str(selected)

        # 手动输入路径
        while True:
            path = input("请输入Shapefile路径: ").strip().strip('"')
            full_path = os.path.join(self.current_dir, path)

            if os.path.exists(full_path) and path.lower().endswith('.shp'):
                print(f"✓ 已选择: {path}")
                print()
                return full_path
            elif os.path.exists(path) and path.lower().endswith('.shp'):
                print(f"✓ 已选择: {path}")
                print()
                return path
            else:
                print("✗ 文件不存在或不是.shp文件，请重新输入")

    def display_shapefile_info(self, reader: ShapefileReader):
        """显示Shapefile信息"""
        print("【步骤 2/7】读取Shapefile信息")
        print(f"✓ 成功读取 {reader.get_record_count()} 条记录")
        print()
        print("字段列表:")
        print("=" * 80)

        field_info = reader.get_field_info()
        for i, field in enumerate(field_info, 1):
            name = field['name']
            dtype = field['type']
            sample = field['sample'] or '(空)'
            # 使用格式化字符串，对齐更清晰
            print(f" [{i:2d}] 字段名: {name:20s} | 类型: {dtype:12s} | 示例: {sample}")

        print("=" * 80)
        print()

    def select_template(self) -> str:
        """选择Word模板"""
        print("【步骤 3/7】选择Word模板文件")
        print()

        # 自动查找.docx文件
        docx_files = list(self.current_dir.glob("*.docx"))
        docx_files = [f for f in docx_files if not f.name.startswith("~$")]  # 排除临时文件

        if docx_files:
            print("找到以下Word模板:")
            for i, docx in enumerate(docx_files, 1):
                print(f"  {i}. {docx.name}")
            print()

            choice = input(f"请选择 (1-{len(docx_files)}, 或按Enter输入路径): ").strip()

            if choice.isdigit() and 1 <= int(choice) <= len(docx_files):
                selected = docx_files[int(choice) - 1]
                print(f"✓ 已选择: {selected.name}")
                print()
                return str(selected)

        # 手动输入路径
        while True:
            path = input("请输入模板路径: ").strip().strip('"')
            full_path = os.path.join(self.current_dir, path)

            if os.path.exists(full_path) and path.lower().endswith('.docx'):
                print(f"✓ 已选择: {path}")
                print()
                return full_path
            elif os.path.exists(path) and path.lower().endswith('.docx'):
                print(f"✓ 已选择: {path}")
                print()
                return path
            else:
                print("✗ 文件不存在或不是.docx文件，请重新输入")

    def display_template_info(self, processor: TemplateProcessor, reader: ShapefileReader):
        """显示模板信息"""
        print("【步骤 4/7】检测模板占位符")
        print()

        placeholders = processor.get_placeholders()

        if not placeholders:
            print("⚠ 警告: 未在模板中检测到任何占位符")
            print("占位符格式应为: !字段名!")
            print()

        print(f"✓ 找到 {len(placeholders)} 个占位符:")
        print("-" * 60)

        fields = reader.get_fields()
        field_set = set(fields)

        for i, placeholder in enumerate(placeholders, 1):
            if placeholder in field_set:
                print(f"  {i}. !{placeholder}! ✓ (匹配字段: {placeholder})")
            else:
                print(f"  {i}. !{placeholder}! ✗ (未找到匹配字段)")

        print("-" * 60)
        print()

    def select_naming_field(self, reader: ShapefileReader) -> str:
        """选择文件命名字段"""
        print("【步骤 5/7】选择文件命名字段")
        print("提示: 该字段的值将用作生成文档的文件名")
        print()

        fields = reader.get_fields()
        field_info = reader.get_field_info()

        # 显示简化的字段列表，方便用户选择
        print("可用字段列表:")
        print("-" * 50)
        for i, field in enumerate(field_info, 1):
            name = field['name']
            dtype = field['type']
            sample = field['sample'] or '(空)'
            # 简化显示，只显示前30个字符的示例值
            sample_short = (sample[:27] + '...') if len(sample) > 30 else sample
            print(f" [{i:2d}] {name:20s} ({dtype:10s}) 示例: {sample_short}")
        print("-" * 50)
        print()

        while True:
            choice = input(f"请输入字段编号 (1-{len(fields)}): ").strip()

            if choice.isdigit() and 1 <= int(choice) <= len(fields):
                idx = int(choice) - 1
                selected_field = fields[idx]
                field_data = field_info[idx]

                # 显示详细的确认信息
                print()
                print("✓ 已选择字段:")
                print(f"   编号: [{int(choice):2d}]")
                print(f"   名称: {field_data['name']}")
                print(f"   类型: {field_data['type']}")
                print(f"   示例: {field_data['sample'] or '(空值)'}")
                print()

                return selected_field
            else:
                print("✗ 无效的选择，请重新输入")

    def select_output_dir(self) -> str:
        """选择输出目录"""
        print("【步骤 6/7】选择输出目录")
        print()

        default_dir = self.current_dir / "output"
        print(f"默认输出目录: {default_dir}")

        user_input = input("按Enter使用默认目录，或输入新路径: ").strip()

        if not user_input:
            print(f"✓ 使用默认目录: output")
            print()
            return str(default_dir)
        else:
            output_path = os.path.join(self.current_dir, user_input)
            print(f"✓ 输出目录: {user_input}")
            print()
            return output_path

    def preview_and_confirm(self, reader: ShapefileReader, naming_field: str, output_dir: str) -> bool:
        """预览并确认"""
        print("【步骤 7/7】预览并确认")
        print("-" * 60)
        print(f"Shapefile记录数: {reader.get_record_count()}")
        print(f"输出目录: {output_dir}")
        print(f"命名字段: {naming_field}")

        # 显示文件名示例
        records = list(reader.get_records())[:min(5, reader.get_record_count())]
        print("文件名示例:")
        for i, record in enumerate(records, 1):
            filename = record.get(naming_field, '')
            if filename:
                print(f"  [{i}] {filename}.docx")
            else:
                print(f"  [{i}] <空值>.docx (字段值为空，将使用默认名称)")

        if reader.get_record_count() > 5:
            print(f"  ... 还有 {reader.get_record_count() - 5} 个文件")

        print("-" * 60)
        print()

        choice = input("确认开始生成? [Y/n]: ").strip().lower()

        if choice in ['', 'y', 'yes']:
            print()
            return True
        else:
            print("操作已取消")
            return False

    def display_results(self, results: Dict[str, Any]):
        """显示生成结果"""
        print()
        print("=" * 60)
        print(" " * 20 + "生成完成")
        print("=" * 60)
        print(f"总计: {results['total']} 个文档")
        print(f"成功: {len(results['success'])} 个")
        print(f"失败: {len(results['failed'])} 个")
        print()

        if results['failed']:
            print("失败列表:")
            for filename, error in results['failed'][:5]:  # 只显示前5个
                print(f"  ✗ {filename}.docx - {error}")
            if len(results['failed']) > 5:
                print(f"  ... 还有 {len(results['failed']) - 5} 个失败")
            print()

        print("=" * 60)


def main():
    """主函数"""
    cli = InteractiveCLI()

    try:
        # 显示欢迎信息
        cli.print_header()

        # 步骤1: 选择Shapefile
        shp_path = cli.select_shapefile()

        # 读取Shapefile
        reader = ShapefileReader(shp_path)

        # 步骤2: 显示Shapefile信息
        cli.display_shapefile_info(reader)

        # 步骤3: 选择模板
        template_path = cli.select_template()

        # 处理模板
        processor = TemplateProcessor(template_path)

        # 步骤4: 显示模板信息
        cli.display_template_info(processor, reader)

        # 步骤5: 选择命名字段
        naming_field = cli.select_naming_field(reader)

        # 步骤6: 选择输出目录
        output_dir = cli.select_output_dir()

        # 步骤7: 预览并确认
        if not cli.preview_and_confirm(reader, naming_field, output_dir):
            sys.exit(0)

        # 批量生成
        generator = BatchGenerator(reader, processor)
        results = generator.generate_all(output_dir, naming_field)

        # 显示结果
        cli.display_results(results)

    except KeyboardInterrupt:
        print("\n\n操作已取消")
        sys.exit(1)
    except Exception as e:
        print(f"\n✗ 错误: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()
